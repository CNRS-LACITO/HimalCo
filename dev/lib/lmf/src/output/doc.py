#! /usr/bin/env python

from config.mdf import mdf_semanticRelation, pd_grammaticalNumber, pd_person, pd_anymacy, pd_clusivity
from utils.error_handling import OutputError
from common.defs import VERNACULAR, ENGLISH, NATIONAL, REGIONAL
from utils.io import ENCODING, EOL

from docx import Document
from docx.shared import Inches

## To define languages and fonts
import config

def doc_write(object, filename, items=lambda lexical_entry: lexical_entry.get_lexeme(), sort_order=None, paradigms=False, reverse=False):
    """! @brief Write a document file.
    @param object The LMF instance to convert into document output format.
    @param filename The name of the document file to write with full path, for instance 'user/output.doc'.
    @param items Lambda function giving the item to sort. Default value is 'lambda lexical_entry: lexical_entry.get_lexeme()', which means that the items to sort are lexemes.
    @param sort_order Python list. Default value is 'None', which means that the document output is alphabetically ordered.
    """
    import string
    if sort_order is None:
        # Lowercase and uppercase letters must have the same rank
        sort_order = dict([(c, ord(c)) for c in string.lowercase])
        up = dict([(c, ord(c) + 32) for c in string.uppercase])
        sort_order.update(up)
        sort_order.update({'':0, ' ':0})
    document = Document()
    # Parse LMF values
    if object.__class__.__name__ == "LexicalResource":
        for lexicon in object.get_lexicons():
            # Document title
            document.add_heading(lexicon.get_id(), 0)
            # Plain paragraph
            document.add_paragraph(lexicon.get_label())
            # Page break
            document.add_page_break()
            # Lexicon is already ordered
            n = -1
            level = 0
            current_item = ("", "")
            previous_character = ''
            current_character = ''
            for lexical_entry in lexicon.get_lexical_entries():
                # Consider only main entries (subentries will be written as parts of the main entry)
                if lexical_entry.find_related_forms("main entry") != []:
                    continue
                if type(sort_order) is type(list()):
                    # Check if item of current element is different from previous one
                    while items(lexical_entry) != current_item[0].decode(ENCODING):
                        try:
                            n += 1
                            current_item = sort_order[n]
                            while current_item[0].startswith("TITLE"):
                                level = int(current_item[0][-1])
                                # Heading
                                document.add_heading(current_item[1].decode(ENCODING), level=level)
                                n += 1
                                current_item = sort_order[n]
                            # Heading
                            document.add_heading(current_item[1].decode(ENCODING), level=level+1)
                            # Paragraph
                            p = document.add_paragraph()
                        except IndexError:
                            # Reached end of list
                            break
                elif type(sort_order) is type(dict()):
                    # Check if current element is a lexeme starting with a different character than previous lexeme
                    try:
                        current_character = items(lexical_entry)[0]
                        if sort_order[items(lexical_entry)[0:1]]:
                            current_character = items(lexical_entry)[0:1]
                        if sort_order[items(lexical_entry)[0:2]]:
                            current_character = items(lexical_entry)[0:2]
                    except IndexError:
                        pass
                    except KeyError:
                        pass
                    except TypeError:
                        pass
                    try:
                        if ( (type(sort_order) is not type(dict())) and ((previous_character == '') or (sort_order(current_character) != sort_order(previous_character))) ) \
                            or ( (type(sort_order) is type(dict())) and (int(sort_order[current_character]) != int(sort_order[previous_character])) ):
                            # Do not consider special characters
                            previous_character = current_character
                            document.add_page_break()
                            title = ''
                            if type(sort_order) is not type(dict()):
                                title += ' ' + current_character
                            else:
                                for key,value in sorted(sort_order.items(), key=lambda x: x[1]):
                                    if int(value) == int(sort_order[current_character]):
                                        title += ' ' + key
                            document.add_heading("-" + title + " -".decode(ENCODING), level=level+1)
                    except KeyError:
                        print Warning("Cannot sort item %s" % items(lexical_entry).encode(ENCODING))
                    except IndexError:
                        # Item is an empty string
                        pass
                if not reverse:
                    # Lexeme
                    lexeme = lexical_entry.get_lexeme()
                    if lexical_entry.get_homonymNumber() is not None:
                        # Add homonym number to lexeme
                        lexeme += " (" + str(lexical_entry.get_homonymNumber()) + ")"
                    # Add dialect if any
                    dialect = ""
                    for sense in lexical_entry.get_senses():
                        for usage_note in sense.find_usage_notes(language=config.xml.vernacular):
                            dialect += " (" + usage_note + ")"
                    p = document.add_paragraph()
                    p.add_run(lexeme).bold = True
                    p.add_run(dialect)
                    # Dialectal variants
                    write_title = True
                    for repr in lexical_entry.get_form_representations():
                        if repr.get_geographicalVariant() is not None:
                            if write_title:
                                p.add_run(" Variante(s) : ")
                                write_title = False
                            else:
                                p.add_run(" ; ")
                            p.add_run(repr.get_geographicalVariant()).bold = True
                            if repr.get_dialect() is not None:
                                p.add_run(" (" + repr.get_dialect() + ")")
                    # Italic
                    if lexical_entry.get_partOfSpeech() is not None:
                        p.add_run(". ")
                        p.add_run(lexical_entry.get_partOfSpeech()).italic = True
                    p.add_run(".")
                    for sense in lexical_entry.get_senses():
                        # Glosses
                        glosses = ""
                        if sense.get_senseNumber() is not None:
                            p = document.add_paragraph()
                            p.add_run("\t" + sense.get_senseNumber() + ")")
                        for gloss in sense.find_glosses(language=config.xml.vernacular):
                            glosses += " " + gloss + " ;"
                        glosses = glosses.rstrip(" ;")
                        if glosses != "":
                            glosses += "."
                        try:
                            for gloss in sense.find_glosses(language=config.xml.French):
                                glosses += " " + gloss + " ;"
                        except AttributeError:
                            for gloss in sense.find_glosses(language=config.xml.English):
                                glosses += " " + gloss + " ;"
                        glosses = glosses.rstrip(" ;")
                        if glosses != "" and glosses[-1] != '.' and glosses[-1] != '!' and glosses[-1] != '?':
                            glosses += "."
                        p.add_run(glosses)
                        # Scientific name
                        if lexical_entry.get_scientific_name() is not None:
                            p.add_run(" ")
                            p.add_run(lexical_entry.get_scientific_name())
                        # Examples
                        for context in sense.get_contexts():
                            p = document.add_paragraph()
                            examples = ""
                            vernacular_forms = context.find_written_forms(language=config.xml.vernacular)
                            for example in vernacular_forms:
                                p.add_run("\t")
                                p.add_run(example).bold = True
                            fra_forms = context.find_written_forms(language=config.xml.French)
                            if len(vernacular_forms) != 0 and len(fra_forms) != 0:
                                p.add_run(" ; ")
                            for example in fra_forms:
                                p.add_run(example)
                            if len(fra_forms) != 0 and fra_forms[0][-1] != '!' and fra_forms[0][-1] != '?':
                                p.add_run(".")
                        # Links
                        if len(lexical_entry.get_related_forms("simple link")) != 0:
                            p = document.add_paragraph()
                            p.add_run("\tVoir :").italic = True
                            for related_form in lexical_entry.get_related_forms("simple link"):
                                if related_form.get_lexical_entry() is not None:
                                    # TODO : hyperlink
                                    pass
                                p.add_run(" ")
                                p.add_run(related_form.get_lexeme()).bold = True
                            p.add_run(".")
                            if len(lexical_entry.find_notes(type="general")) != 0:
                                p.add_run(" ")
                        # Notes
                        if len(lexical_entry.find_notes(type="general")) != 0:
                            if len(lexical_entry.get_related_forms("simple link")) == 0:
                                p = document.add_paragraph()
                                p.add_run("\t")
                            p.add_run("[Note :").italic = True
                            for note in lexical_entry.find_notes(type="general"):
                                p.add_run(" ")
                                p.add_run(note)
                            p.add_run("].").italic = True
                    if paradigms:
                        # Intense quote
                        document.add_paragraph('Paradigms', style='IntenseQuote')
                        # Table
                        table = document.add_table(rows=1, cols=2)
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = 'Paradigm'
                        hdr_cells[1].text = 'Form'
                        for item in lexical_entry.find_paradigms(grammatical_number=pd_grammaticalNumber["sg"]):
                            row_cells = table.add_row().cells
                            row_cells[0].text = "sg"
                            row_cells[1].text = item
                        for item in lexical_entry.find_paradigms(grammatical_number=pd_grammaticalNumber["pl"]):
                            row_cells = table.add_row().cells
                            row_cells[0].text = "pl"
                            row_cells[1].text = item
                        for item in lexical_entry.find_paradigms(person=pd_person[1], grammatical_number=pd_grammaticalNumber['s']):
                            row_cells = table.add_row().cells
                            row_cells[0].text = "1s"
                            row_cells[1].text = item
                        for item in lexical_entry.find_paradigms(person=pd_person[2], grammatical_number=pd_grammaticalNumber['s']):
                            row_cells = table.add_row().cells
                            row_cells[0].text = "2s"
                            row_cells[1].text = item
                        for item in lexical_entry.find_paradigms(person=pd_person[3], grammatical_number=pd_grammaticalNumber['s']):
                            row_cells = table.add_row().cells
                            row_cells[0].text = "3s"
                            row_cells[1].text = item
                        for item in lexical_entry.find_paradigms(anymacy=pd_anymacy[4], grammatical_number=pd_grammaticalNumber['s']):
                            row_cells = table.add_row().cells
                            row_cells[0].text = "4s"
                            row_cells[1].text = item
                        for item in lexical_entry.find_paradigms(person=pd_person[1], grammatical_number=pd_grammaticalNumber['d']):
                            row_cells = table.add_row().cells
                            row_cells[0].text = "1d"
                            row_cells[1].text = item
                        for item in lexical_entry.find_paradigms(person=pd_person[2], grammatical_number=pd_grammaticalNumber['d']):
                            row_cells = table.add_row().cells
                            row_cells[0].text = "2d"
                            row_cells[1].text = item
                        for item in lexical_entry.find_paradigms(person=pd_person[3], grammatical_number=pd_grammaticalNumber['d']):
                            row_cells = table.add_row().cells
                            row_cells[0].text = "3d"
                            row_cells[1].text = item
                        for item in lexical_entry.find_paradigms(anymacy=pd_anymacy[4], grammatical_number=pd_grammaticalNumber['d']):
                            row_cells = table.add_row().cells
                            row_cells[0].text = "4d"
                            row_cells[1].text = item
                        for item in lexical_entry.find_paradigms(person=pd_person[1], grammatical_number=pd_grammaticalNumber['p']):
                            row_cells = table.add_row().cells
                            row_cells[0].text = "1p"
                            row_cells[1].text = item
                        for item in lexical_entry.find_paradigms(person=pd_person[1], grammatical_number=pd_grammaticalNumber['p'], clusivity=pd_clusivity['e']):
                            row_cells = table.add_row().cells
                            row_cells[0].text = "1e"
                            row_cells[1].text = item
                        for item in lexical_entry.find_paradigms(person=pd_person[1], grammatical_number=pd_grammaticalNumber['p'], clusivity=pd_clusivity['i']):
                            row_cells = table.add_row().cells
                            row_cells[0].text = "1i"
                            row_cells[1].text = item
                        for item in lexical_entry.find_paradigms(person=pd_person[2], grammatical_number=pd_grammaticalNumber['p']):
                            row_cells = table.add_row().cells
                            row_cells[0].text = "2p"
                            row_cells[1].text = item
                        for item in lexical_entry.find_paradigms(person=pd_person[3], grammatical_number=pd_grammaticalNumber['p']):
                            row_cells = table.add_row().cells
                            row_cells[0].text = "3p"
                            row_cells[1].text = item
                        for item in lexical_entry.find_paradigms(anymacy=pd_anymacy[4], grammatical_number=pd_grammaticalNumber['p']):
                            row_cells = table.add_row().cells
                            row_cells[0].text = "4p"
                            row_cells[1].text = item
                        # Paragraph
                        p = document.add_paragraph()
                        # Italic
                        p.add_run(lexical_entry.find_notes(type=None)).italic = True
                        # Picture
                        #document.add_picture('name.png', width=Inches(1.25))
                    # Handle subentries
                    for related_form in lexical_entry.get_related_forms("subentry"):
                        if related_form.get_lexical_entry() is not None:
                            # Items in unordered list
                            p = document.add_paragraph(style='ListBullet')
                            p.add_run(related_form.get_lexeme()).bold = True
                            for sense in related_form.get_lexical_entry().get_senses():
                                glosses = ""
                                for gloss in sense.find_glosses(language=config.xml.vernacular):
                                    glosses += " " + gloss + " ;"
                                glosses = glosses.rstrip(" ;")
                                if glosses != "":
                                    glosses += "."
                                try:
                                    for gloss in sense.find_glosses(language=config.xml.French):
                                        glosses += " " + gloss + " ;"
                                except AttributeError:
                                    for gloss in sense.find_glosses(language=config.xml.English):
                                        glosses += " " + gloss + " ;"
                                glosses = glosses.rstrip(" ;")
                                if glosses != "":
                                    glosses += "."
                                p.add_run(glosses)
                    p.add_run(EOL)
                else: # reverse
                    # Paragraph
                    p = document.add_paragraph()
                    # French gloss
                    for sense in lexical_entry.get_senses():
                        for gloss in sense.find_glosses(language=config.xml.French):
                            p.add_run(gloss).bold = True
                            break
                    # Scientific name
                    if lexical_entry.get_scientific_name() is not None:
                        p.add_run(" ")
                        p.add_run(lexical_entry.get_scientific_name())
                    p.add_run(". ")
                    # Lexeme
                    p.add_run(lexical_entry.get_lexeme())
                    if lexical_entry.get_lexeme()[-1] != '?' and lexical_entry.get_lexeme()[-1] != '!' and lexical_entry.get_lexeme()[-1] != '.':
                        p.add_run(".")
    else:
        raise OutputError(object, "Object to write must be a Lexical Resource.")
    document.save(filename)
